"""
Text extraction service for PDF and DOCX files.
Handles conversion of CV documents to plain text.
"""

import os
from typing import Optional
import PyPDF2
from docx import Document


class TextExtractor:
    """Extract text from PDF and DOCX files."""
    
    def extract(self, file_path: str) -> Optional[str]:
        """
        Extract text from a PDF or DOCX file.
        
        Args:
            file_path (str): Path to the file (PDF or DOCX)
            
        Returns:
            Optional[str]: Extracted text, None if extraction fails
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.pdf':
            return self._extract_from_pdf(file_path)
        elif file_ext == '.docx':
            return self._extract_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """
        Extract text from PDF file.
        
        Args:
            file_path (str): Path to PDF file
            
        Returns:
            str: Extracted text from all pages
        """
        text = ""
        
        try:
            with open(file_path, 'rb') as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n"
        
        except Exception as e:
            raise ValueError(f"Error extracting PDF: {str(e)}")
        
        return text.strip()
    
    def _extract_from_docx(self, file_path: str) -> str:
        """
        Extract text from DOCX file.
        
        Args:
            file_path (str): Path to DOCX file
            
        Returns:
            str: Extracted text from all paragraphs
        """
        text = ""
        
        try:
            doc = Document(file_path)
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Also extract text from tables if present
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " "
                    text += "\n"
        
        except Exception as e:
            raise ValueError(f"Error extracting DOCX: {str(e)}")
        
        return text.strip()
    
    def clean_text(self, text: str) -> str:
        """
        Clean extracted text by removing extra whitespace and normalizing.
        
        Args:
            text (str): Raw extracted text
            
        Returns:
            str: Cleaned text
        """
        # Remove multiple spaces
        text = " ".join(text.split())
        
        # Normalize line breaks
        lines = text.split("\n")
        lines = [line.strip() for line in lines if line.strip()]
        
        return "\n".join(lines)
